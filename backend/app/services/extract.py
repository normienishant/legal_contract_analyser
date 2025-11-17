"""Document extraction service."""
import os
import re
from typing import List
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats."""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from document based on extension."""
        ext = Path(file_path).suffix.lower()
        
        if ext == ".pdf":
            return DocumentExtractor._extract_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentExtractor._extract_docx(file_path)
        elif ext == ".txt":
            return DocumentExtractor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            doc.close()
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF. The file may be image-based or corrupted.")
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            return extracted_text
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            if not paragraphs:
                raise ValueError("No text could be extracted from DOCX. The file may be empty or corrupted.")
            
            extracted_text = "\n".join(paragraphs)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def _extract_txt(file_path: str) -> str:
        """Extract text from TXT with multiple encoding support."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error extracting TXT with {encoding}: {e}")
                if encoding == encodings[-1]:  # Last encoding failed
                    raise
        
        # If all encodings fail, try with error handling
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                logger.warning("Reading TXT file with error replacement (some characters may be lost)")
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            raise ValueError(f"Unable to read text file: {str(e)}")
    
    @staticmethod
    def segment_clauses(text: str) -> List[str]:
        """
        Robust clause segmentation using regex pattern matching on entire text.
        
        Handles:
        - Articles (ARTICLE—1, ARTICLE 1, Article 1, etc.)
        - Decimal numbered clauses (1.0, 1.1, 1.2, 2.0, etc.)
        - Numbered clauses (1., 2., 10., 1), etc.)
        - Lettered sub-clauses (a), b), a., b., etc.)
        - Multi-line clauses properly
        """
        if not text or not text.strip():
            return []
        
        # Normalize text: preserve newlines but normalize multiple spaces
        # Replace multiple newlines with single newline, but keep single newlines
        text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
        text = re.sub(r'\r', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
        
        clauses = []
        
        # METHOD 1: Split by ARTICLE markers (major sections)
        # Pattern matches: "ARTICLE—1", "ARTICLE 1", "Article 1", "ARTICLE-1", etc.
        article_pattern = r'(?i)(ARTICLE[—\-\s]+\d+)'
        article_matches = list(re.finditer(article_pattern, text))
        
        if len(article_matches) > 0:
            # We found articles, process each section
            last_end = 0
            
            for i, match in enumerate(article_matches):
                start_pos = match.start()
                article_text = match.group(0)
                
                # Process text before this article (if any)
                if start_pos > last_end:
                    prefix_text = text[last_end:start_pos].strip()
                    if prefix_text:
                        section_clauses = DocumentExtractor._segment_section(prefix_text)
                        clauses.extend(section_clauses)
                
                # Get content from this article to next article (or end)
                if i + 1 < len(article_matches):
                    end_pos = article_matches[i + 1].start()
                    article_content = text[match.end():end_pos]
                else:
                    article_content = text[match.end():]
                
                # Combine article header with its content
                full_section = article_text + article_content
                section_clauses = DocumentExtractor._segment_section(full_section)
                clauses.extend(section_clauses)
                
                last_end = match.end() + len(article_content)
            
            # Handle any remaining text after last article
            if last_end < len(text):
                remaining_text = text[last_end:].strip()
                if remaining_text:
                    section_clauses = DocumentExtractor._segment_section(remaining_text)
                    clauses.extend(section_clauses)
        else:
            # No articles found, process entire text as one section
            clauses = DocumentExtractor._segment_section(text)
        
        # Post-process: Clean and validate clauses
        cleaned_clauses = []
        for clause in clauses:
            # Normalize whitespace (multiple spaces/tabs to single space)
            clause = re.sub(r'[ \t]+', ' ', clause.strip())
            # Remove excessive newlines within clause
            clause = re.sub(r'\n+', ' ', clause)
            
            # Filter by length
            if len(clause) >= 15:  # Minimum clause length
                # Truncate if too long (but log warning)
                if len(clause) > 10000:
                    logger.warning(f"Clause truncated from {len(clause)} to 10000 chars")
                    clause = clause[:10000] + "..."
                cleaned_clauses.append(clause)
        
        if cleaned_clauses:
            logger.info(f"Segmented text into {len(cleaned_clauses)} clauses")
            return cleaned_clauses
        
        # FALLBACK: If no clauses found, try paragraph-based splitting
        logger.warning("Primary segmentation failed, using paragraph fallback")
        return DocumentExtractor._fallback_segment(text)
    
    @staticmethod
    def _segment_section(section_text: str) -> List[str]:
        """
        Segment a section (article or entire document) into clauses.
        Handles decimal numbering (1.0, 1.1), regular numbering (1., 2.), etc.
        """
        clauses = []
        
        # METHOD 1A: Split by decimal numbered clauses (1.0, 1.1, 2.0, etc.)
        # Pattern: number.number followed by space, at start of line or after newline
        # Use finditer to get all matches with positions
        decimal_pattern = r'(?:^|\n)\s*(\d+\.\d+\s+)'
        matches = list(re.finditer(decimal_pattern, section_text, re.MULTILINE))
        
        if len(matches) > 0:
            # Found decimal numbered clauses
            last_end = 0
            
            for i, match in enumerate(matches):
                start_pos = match.start()
                clause_num = match.group(1)
                
                # Get content from last match end (or start) to current match start
                if start_pos > last_end:
                    # There's content before this clause
                    prefix = section_text[last_end:start_pos].strip()
                    if prefix and len(prefix) > 15:
                        clauses.append(prefix)
                
                # Get content from current match to next match (or end)
                if i + 1 < len(matches):
                    end_pos = matches[i + 1].start()
                    clause_content = section_text[match.end():end_pos].strip()
                else:
                    clause_content = section_text[match.end():].strip()
                
                # Combine clause number with content
                full_clause = clause_num + clause_content
                
                # Check if this clause contains nested decimal sub-clauses on new lines
                # If so, split them
                if re.search(r'\n\s+\d+\.\d+\s+', clause_content):
                    # Contains nested decimal clauses, split them
                    nested_splits = re.split(r'(?=\n\s+\d+\.\d+\s+)', clause_content)
                    for nested_part in nested_splits:
                        nested_part = nested_part.strip()
                        if nested_part:
                            if re.match(r'\d+\.\d+\s+', nested_part):
                                # This nested part starts with its own decimal number
                                clauses.append(nested_part)
                            else:
                                # This is continuation of main clause
                                clauses.append(clause_num + nested_part)
                else:
                    # Single clause without nested decimals
                    clauses.append(full_clause)
                
                # Update last_end to the end of current clause content
                if i + 1 < len(matches):
                    last_end = matches[i + 1].start()
                else:
                    last_end = len(section_text)
            
            # Handle any remaining text after last match
            if last_end < len(section_text):
                remaining = section_text[last_end:].strip()
                if remaining and len(remaining) > 15:
                    clauses.append(remaining)
        else:
            # No decimal numbering found, try regular numbered clauses
            clauses = DocumentExtractor._segment_by_numbered_clauses(section_text)
        
        return clauses
    
    @staticmethod
    def _segment_by_numbered_clauses(text: str) -> List[str]:
        """
        Segment text by regular numbered clauses (1., 2., 1), etc.)
        """
        clauses = []
        
        # Pattern: number followed by . or ) and space, at start of line or after newline
        # Exclude decimal numbers (already handled) - use negative lookahead
        numbered_pattern = r'(?:^|\n)\s*(\d+[\.\)]\s+)(?!\d)'
        matches = list(re.finditer(numbered_pattern, text, re.MULTILINE))
        
        if len(matches) > 0:
            last_end = 0
            
            for i, match in enumerate(matches):
                start_pos = match.start()
                clause_num = match.group(1)
                
                # Process text before this clause (if any)
                if start_pos > last_end:
                    prefix = text[last_end:start_pos].strip()
                    if prefix and len(prefix) > 15:
                        clauses.append(prefix)
                
                # Get content from current match to next match (or end)
                if i + 1 < len(matches):
                    end_pos = matches[i + 1].start()
                    clause_content = text[match.end():end_pos].strip()
                else:
                    clause_content = text[match.end():].strip()
                
                # Combine clause number with content
                full_clause = clause_num + clause_content
                
                # Check for lettered sub-clauses (a), b), etc.) within
                # Only split if clause is very long (>2000 chars)
                if len(clause_content) > 2000 and re.search(r'\n\s+[a-z][\.\)]\s+', clause_content, re.IGNORECASE):
                    # Try to split by lettered sub-clauses
                    lettered_splits = re.split(r'(?=\n\s+[a-z][\.\)]\s+)', clause_content, flags=re.IGNORECASE)
                    if len(lettered_splits) > 1:
                        for lettered_part in lettered_splits:
                            lettered_part = lettered_part.strip()
                            if lettered_part:
                                if re.match(r'[a-z][\.\)]\s+', lettered_part, re.IGNORECASE):
                                    # Lettered sub-clause, keep with main clause number prefix
                                    clauses.append(clause_num + lettered_part)
                                else:
                                    # Continuation text
                                    clauses.append(clause_num + lettered_part)
                    else:
                        clauses.append(full_clause)
                else:
                    # Single clause without splitting
                    clauses.append(full_clause)
                
                # Update last_end
                if i + 1 < len(matches):
                    last_end = matches[i + 1].start()
                else:
                    last_end = len(text)
            
            # Handle any remaining text after last match
            if last_end < len(text):
                remaining = text[last_end:].strip()
                if remaining and len(remaining) > 15:
                    clauses.append(remaining)
        else:
            # No numbered clauses, try paragraph-based
            clauses = DocumentExtractor._segment_by_paragraphs(text)
        
        return clauses
    
    @staticmethod
    def _segment_by_paragraphs(text: str) -> List[str]:
        """
        Segment text by paragraphs with clause markers.
        """
        clauses = []
        
        # Split by double newlines (paragraph breaks)
        paragraphs = re.split(r'\n\s*\n+', text)
        
        clause_markers = [
            r'^ARTICLE[—\-\s]+\d+',
            r'^\d+\.\d+\s+',
            r'^\d+[\.\)]\s+',
            r'^[a-z][\.\)]\s+',
            r'^WHEREAS\s+',
            r'^THEREFORE\s+',
            r'^NOW\s+THEREFORE\s+',
            r'^IN\s+CONSIDERATION\s+',
            r'^THE\s+PARTIES\s+AGREE\s+',
            r'^Article\s+\d+',
            r'^Section\s+\d+',
            r'^Clause\s+\d+',
        ]
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Normalize whitespace
            para = re.sub(r'\s+', ' ', para)
            
            # Check if paragraph starts with clause marker
            is_clause_start = any(re.match(marker, para, re.IGNORECASE) for marker in clause_markers)
            
            if is_clause_start or len(para) > 50:
                # Check for clause indicators
                if is_clause_start or re.search(
                    r'\b(shall|must|will|agrees?|warrants?|represents?|agreement|contract|tenant|landlord|party|parties|service|provider|institute)\b',
                    para, re.IGNORECASE
                ):
                    clauses.append(para)
        
        return clauses
    
    @staticmethod
    def _fallback_segment(text: str) -> List[str]:
        """
        Final fallback: split by sentences if no other method worked.
        """
        # Normalize text
        normalized_text = re.sub(r'\s+', ' ', text.strip())
        
        # Split by sentence endings
        sentences = re.split(r'([.!?]+\s+)', normalized_text)
        
        clauses = []
        current_clause = ""
        
        for i in range(0, len(sentences), 2):
            sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else "")
            sentence = sentence.strip()
            
            if sentence:
                current_clause += " " + sentence if current_clause else sentence
                
                # Group 2-3 sentences per clause
                if len(current_clause) > 200 or (i + 2 >= len(sentences)):
                    if len(current_clause.strip()) > 15:
                        clauses.append(current_clause.strip())
                    current_clause = ""
        
        if current_clause.strip() and len(current_clause.strip()) > 15:
            clauses.append(current_clause.strip())
        
        return clauses

